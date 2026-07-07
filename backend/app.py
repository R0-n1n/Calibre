"""
backend/app.py
--------------
FastAPI backend for ResumeIQ.

Handles HTTP routing and file parsing only.
All AI work is delegated to utils/ai_provider.py.
This file does NOT import any AI SDK.
"""

import io
import json
import os
from pathlib import Path

import docx
import pdfplumber
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles

from utils.ai_provider import (
    generate_ai_feedback,
    generate_cover_letter,
    generate_interview_questions,
    rewrite_resume_section,
)

load_dotenv(dotenv_path=Path(__file__).parent / ".env")

app = FastAPI(title="Calibre API", version="3.0.0")

origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:5173").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── File parsers ───────────────────────────────────────────────────────────────

def extract_pdf_text(data: bytes) -> str:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        pages = [p.extract_text() or "" for p in pdf.pages]
    text = "\n".join(pages).strip()
    if not text:
        raise ValueError("No text found in PDF. Use a text-based PDF, not a scanned image.")
    return text


def extract_docx_text(data: bytes) -> str:
    doc = docx.Document(io.BytesIO(data))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in parts:
                    parts.append(t)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("No text found in DOCX file.")
    return text


def parse_resume(upload: UploadFile, data: bytes) -> str:
    name = (upload.filename or "").lower()
    ctype = (upload.content_type or "").lower()
    if name.endswith(".pdf") or "pdf" in ctype:
        return extract_pdf_text(data)
    elif name.endswith(".docx") or "wordprocessingml" in ctype:
        return extract_docx_text(data)
    else:
        raise ValueError(f"Unsupported file: '{upload.filename}'. Upload a PDF or DOCX.")


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {"status": "ok", "provider": os.getenv("AI_PROVIDER", "groq")}


@app.post("/api/analyze")
async def analyze(
    resume: UploadFile = File(...),
    job_description: str = Form(...),
):
    try:
        data = await resume.read()
        resume_text = parse_resume(resume, data)
        result = generate_ai_feedback(resume_text, job_description)
        result["resume_text"] = resume_text  # cached by frontend for reuse
        return JSONResponse(content=result)
    except ValueError as e:
        raise HTTPException(422, str(e))
    except KeyError as e:
        raise HTTPException(500, f"Missing env variable: {e}. Check backend/.env")
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/rewrite")
async def rewrite(
    section: str = Form(...),
    content: str = Form(...),
    job_description: str = Form(...),
):
    try:
        return {"rewritten": rewrite_resume_section(section, content, job_description), "section": section}
    except KeyError as e:
        raise HTTPException(500, f"Missing env variable: {e}. Check backend/.env")
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/cover-letter")
async def cover_letter(
    resume_text: str = Form(...),
    job_description: str = Form(...),
    analysis: str = Form(...),
):
    try:
        return {"cover_letter": generate_cover_letter(resume_text, job_description, json.loads(analysis))}
    except json.JSONDecodeError:
        raise HTTPException(422, "Invalid analysis JSON.")
    except KeyError as e:
        raise HTTPException(500, f"Missing env variable: {e}. Check backend/.env")
    except Exception as e:
        raise HTTPException(500, str(e))


@app.post("/api/interview-questions")
async def interview_questions(
    resume: UploadFile = File(...),
    job_description: str = Form(...),
):
    try:
        data = await resume.read()
        resume_text = parse_resume(resume, data)
        return {"questions": generate_interview_questions(resume_text, job_description)}
    except ValueError as e:
        raise HTTPException(422, str(e))
    except Exception as e:
        raise HTTPException(500, str(e))


# ── Serve built frontend in production ─────────────────────────────────────────
_dist = Path(__file__).parent.parent / "frontend" / "dist"
if _dist.exists():
    _assets = _dist / "assets"
    if _assets.exists():
        app.mount("/assets", StaticFiles(directory=str(_assets)), name="assets")

    @app.get("/{full_path:path}")
    async def serve_spa(full_path: str):
        index = _dist / "index.html"
        return FileResponse(str(index)) if index.exists() else JSONResponse(
            {"detail": "Frontend not built. Run: cd frontend && npm run build"}, 404
        )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True, reload_dirs=[".", "utils"])
